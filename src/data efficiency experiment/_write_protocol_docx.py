"""Generate PROTOCOL.docx — human-readable protocol outside the IDE."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parent / "PROTOCOL.docx"


def _set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    for i in range(1, 4):
        h = doc.styles[f"Heading {i}"]
        h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        h.font.name = "Calibri"


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.rows[i + 1].cells[j].text = val
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    _set_doc_defaults(doc)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    title = doc.add_heading(
        "Data-Efficiency Experiment — Training & Evaluation Protocol", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph(
        "This document is the single source of truth for training and evaluating "
        "plain (λ = 0) and physics (λ > 0) Multires models across training-set sizes. "
        "Follow it as written; do not re-negotiate steps ad hoc."
    )
    doc.add_paragraph(
        "Code in the experiment folder may import existing modules from the parent "
        "src directory (e.g. data_generation, dataset_utils, data_c_amb_loss_diagnostics, "
        "train, frognet). From notebooks or scripts in this folder, run:"
    )
    p = doc.add_paragraph()
    r = p.add_run("import setup_src_path  # adds parent src/ to sys.path")
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    doc.add_paragraph(
        "Or keep the working directory as src. Prefer writing artifacts under this "
        "folder (checkpoints/, runs/) with paths anchored to this directory."
    )

    doc.add_heading("1. Experiment scope", level=1)
    doc.add_paragraph(
        "For each training size n_train below, train and evaluate:"
    )
    add_numbered(
        doc,
        [
            "Plain network: λ = 0",
            "Physics network: λ > 0 (TRACE loss term), with trace_scale = 8 always",
        ],
    )
    doc.add_heading("Training-set sizes", level=2)
    doc.add_paragraph(
        "n_train ∈ {300, 866, 2498, 7207, 20794, 60000}"
    )
    doc.add_paragraph(
        "Process sizes from small to large. For each n: plain (LR tune → final) → "
        "physics (λ tune → final) → shared test."
    )
    doc.add_heading("Data family (all models)", level=2)
    doc.add_paragraph(
        "Every plain and physics model — for every n_train — is trained and evaluated "
        "on the same pulse/TRACE family: spectrally filtered C1, matching "
        "filtered_c1_multires_2k_diagnostics_NB.ipynb (FilteredC1PulseConfig / "
        "generate_pulses_filtered_c1 + physical FROGNet). Do not mix Data C / "
        "unfiltered C1 / other generators into this campaign."
    )

    doc.add_heading("2. Shared fixed settings (all models, all n)", level=1)
    add_table(
        doc,
        ["Item", "Value"],
        [
            ["Architecture", "Multires"],
            [
                "Pulse / TRACE data",
                "Filtered C1 ONLY (same as filtered_c1_multires_2k_diagnostics_NB.ipynb); "
                "train, val, and test for all models",
            ],
            ["Batch size B", "300"],
            ["drop_last", "False"],
            ["Optimizer", "Adam"],
            [
                "n_val",
                "200 — identical val set for every model and every n",
            ],
            [
                "n_test",
                "Fixed (e.g. 512) — identical test set for every model and every n",
            ],
            [
                "Train / val SNR",
                "Uniform in [0, 30] dB (resampled each use, as in the diagnostics notebook)",
            ],
            [
                "Selection metric",
                "Validation pulse L1 after best ambiguity",
            ],
            [
                "K (steps/epoch)",
                "K = ceil(n_train / B)",
            ],
            [
                "Early stopping",
                "Validate every epoch (curves vs epoch). Stop when "
                "(current_step - t_best) >= patience_steps with "
                "patience_steps = 25 * K (= 25 epochs). Same for screens and finals.",
            ],
            [
                "Step budget",
                "Screens: max_steps = 100 * K (= 100 epochs). "
                "Finals: max_steps = 200 * K (= 200 epochs). "
                "May set max_epochs = ceil(max_steps / K).",
            ],
            [
                "Physics TRACE scale",
                "trace_scale = 8 for all physics runs (do not use a data-dependent scale)",
            ],
            ["Canonicalization", "Same as diagnostics notebook (t0)"],
            [
                "Seeds",
                "Train pulses may depend on n; val and test pulses/seeds fixed once and reused for all models",
            ],
        ],
    )
    doc.add_paragraph("Budget table (B = 300), all campaign n_train:")
    add_table(
        doc,
        [
            "N",
            "K",
            "patience (steps)",
            "patience (epochs)",
            "max screen (steps)",
            "screen (epochs)",
            "max final (steps)",
            "final (epochs)",
        ],
        [
            ["300", "1", "25", "25", "100", "100", "200", "200"],
            ["866", "3", "75", "25", "300", "100", "600", "200"],
            ["2498", "9", "225", "25", "900", "100", "1800", "200"],
            ["7207", "25", "625", "25", "2500", "100", "5000", "200"],
            ["20794", "70", "1750", "25", "7000", "100", "14000", "200"],
            ["60000", "200", "5000", "25", "20000", "100", "40000", "200"],
        ],
    )

    doc.add_heading("3. Logging during training (plain and physics)", level=1)
    doc.add_paragraph(
        "Every epoch, record (same spirit as Multires 2K plain/physics in "
        "filtered_c1_multires_2k_diagnostics_NB.ipynb):"
    )
    add_numbered(
        doc,
        [
            "Pulse L1 (train & val): raw and best-ambiguity",
            "TRACE L1 (train & val)",
            "Gradient norms: ||∇L_data||, ||∇L_reg||, ||∇L_total||. "
            "For λ = 0: total loss is L_data only; still compute ||∇L_reg|| for diagnostics "
            "and label it as not entering the loss.",
            "Per-batch timings (ms/batch mean ± over epochs): "
            "data_prep_sec, loss_data_fwd_sec, loss_reg_fwd_sec, total_backward_sec, optimizer_step_sec",
            "Wall-clock: dataset generation time; total train time; time to best checkpoint",
        ],
    )
    doc.add_heading("Reconstruction evolution snapshots", level=2)
    doc.add_paragraph(
        "Every 2 epochs (and also at best / stop if those epochs are odd):"
    )
    add_bullets(
        doc,
        [
            "Fix one validation example for the whole experiment (same sample for all models).",
            "Apply a fixed noisy TRACE at SNR = 10 dB (same noise realization across epochs).",
            "Save I_noisy, E_true, E_pred (+ epoch / step).",
        ],
    )
    doc.add_paragraph("Do not store full-set predictions each epoch.")

    doc.add_heading("4. Plain network (λ = 0) — LR tuning and final train", level=1)
    doc.add_paragraph(
        "Goal: choose learning rate, then produce the official plain checkpoint for this n."
    )
    doc.add_heading("4.1 LR grid (initial)", level=2)
    doc.add_paragraph("Scan exactly:")
    doc.add_paragraph("LR ∈ {10⁻⁴, 10⁻³, 10⁻²}")
    doc.add_heading("4.2 If LR* is at a grid edge — expand", level=2)
    doc.add_paragraph(
        "Do not accept an edge winner without at least one expansion toward that edge."
    )
    add_bullets(
        doc,
        [
            "Lower edge (LR* = 10⁻⁴): add {3·10⁻⁵, 10⁻⁵}, re-screen, update LR*.",
            "Upper edge (LR* = 10⁻²): add {3·10⁻², 10⁻¹}, re-screen, update LR*.",
            "Repeat at most 1–2 expansions until LR* is interior or the expanded edge remains best.",
            "Optional fine point: geometric midpoint between LR* and its better neighbor "
            "(e.g. between 10⁻⁴ and 10⁻³ try ~3·10⁻⁴).",
        ],
    )
    doc.add_heading("4.3 Procedure", level=2)
    add_numbered(
        doc,
        [
            "For each LR in the current grid, train with §2 settings and §3 logging; "
            "early stop = fixed step patience (§2).",
            "Select LR* by best validation pulse L1 (best-amb). If at an edge, apply §4.2 "
            "before declaring LR*.",
            "Final plain training with LR*, same early-stop rule, max_steps, full logging "
            "and snapshots.",
            "Save: best checkpoint, history (curves + grads + timings), snapshots, "
            "meta (n, LR*, λ = 0, seeds, best/stop step).",
        ],
    )
    doc.add_paragraph(
        "There is no TRACE weight in the optimized loss (λ = 0)."
    )

    doc.add_heading("5. Physics network (λ > 0) — λ tuning and final train", level=1)
    doc.add_paragraph(
        "Goal: choose λ with the plain network’s LR*; produce the official physics "
        "checkpoint for this n."
    )
    doc.add_heading("5.1 Fixed for all physics runs", level=2)
    add_bullets(
        doc,
        [
            "trace_scale = 8",
            "LR = LR* from the plain model at the same n",
            "Do not retune LR after finding λ*",
        ],
    )
    doc.add_paragraph(
        "Loss form (conceptually): L = L_data + λ · L_reg, where L_reg uses TRACE scale 8."
    )
    doc.add_heading("5.2 Coarse λ search (short training)", level=2)
    doc.add_paragraph("Log-spaced candidates over [0.75, 15], e.g.:")
    doc.add_paragraph("λ ∈ {0.75, 1.5, 3, 6, 12}")
    doc.add_paragraph("(optional extra point 15).")
    doc.add_paragraph(
        "Prior for small data: if n_train ≲ 2000 (here: 300, 866; treat 2498 conservatively "
        "the same if desired), do not favor λ < 0.75; start the coarse grid at λ ≥ 3 "
        "(e.g. {3, 6, 12, 15})."
    )
    doc.add_paragraph(
        "Use the same fixed step-patience early-stop rule as finals (section 2); optional "
        "explicit shorter screen rule only if labeled. Rank by validation pulse L1 (best-amb)."
    )
    doc.add_heading("5.3 Fine λ search", level=2)
    doc.add_paragraph(
        "Around the coarse winner λ*, local geometric bisection (log-λ):"
    )
    add_bullets(
        doc,
        [
            "Evaluate geometric midpoints between λ* and its coarse neighbors "
            "(e.g. if 6 wins between 3 and 12, try √(3·6) and √(6·12)).",
            "Optionally one more bisection (2–4 fine points total).",
            "Keep screening short; pick the best λ on val.",
        ],
    )
    doc.add_heading("5.4 Final physics training", level=2)
    doc.add_paragraph(
        "One full run with (λ*, LR*), trace_scale = 8, early-stop = fixed step patience "
        "(section 2), max_steps, full logging and snapshots. Save checkpoint + history + meta."
    )

    doc.add_heading("6. Evaluation (after each final checkpoint)", level=1)
    doc.add_paragraph("Use the same held-out test set for every model:")
    add_numbered(
        doc,
        [
            "SNR sweep (same grid as the diagnostics notebook, e.g. [−10, 30] dB step 5): "
            "report best-amb pulse L1 and similarity (and PCGPA if included in the campaign).",
            "Single test example @ SNR = 10 dB: clean TRACE, noisy TRACE, |E(t)|, phase after "
            "best-amb — for plain and physics.",
            "Summary row per model: n_train, plain/physics, LR*, λ (0 or λ*), "
            "trace_scale (8 if physics), best/stop step, key val/test metrics, "
            "wall-clock and mean batch timings.",
        ],
    )

    doc.add_heading("7. Per-n checklist", level=1)
    doc.add_paragraph("For a given n_train:")
    add_bullets(
        doc,
        [
            "Build / load shared val (200) and test; build train of size n",
            "Plain: LR grid {10⁻⁴, 10⁻³, 10⁻²}; expand if LR* at edge (4.2); then final plain train + logs + snapshots",
            "Physics: LR* fixed, trace_scale = 8; coarse λ → fine λ → final physics train + logs + snapshots",
            "Test SNR sweep + example @ 10 dB on both finals",
            "Archive meta and timings for cross-n comparison",
        ],
    )

    doc.add_heading("8. Fairness notes (do not violate)", level=1)
    add_bullets(
        doc,
        [
            "Same B, val set, test set, early-stop rule (fixed step patience), max_steps ceiling, SNR policy, and architecture across models.",
            "Different K (steps/epoch) across n is expected; do not force equal step counts to convergence.",
            "Tuning compute (LR / λ screens) should be logged separately from final-train wall time when comparing efficiency.",
            "Physics always uses trace_scale = 8.",
            "No LR retuning for physics after λ* is chosen.",
            "Do not declare LR* at a grid edge without the expansion in 4.2.",
        ],
    )

    doc.add_heading("9. Suggested artifact layout (under this folder)", level=1)
    layout = doc.add_paragraph()
    run = layout.add_run(
        "data efficiency experiment/\n"
        "  PROTOCOL.docx              # this file (readable outside the IDE)\n"
        "  PROTOCOL.md                # markdown mirror (optional)\n"
        "  setup_src_path.py          # import helper for parent src/\n"
        "  checkpoints/\n"
        "    n{n}_plain_lr{lr}/...\n"
        "    n{n}_phys_lam{lam}_lr{lr}/...\n"
        "  sweeps/                    # optional LR / lambda screen summaries\n"
        "  figures/                   # optional plots"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    doc.add_paragraph(
        "Use consistent tags including n, lam, lr, and seed in filenames."
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
